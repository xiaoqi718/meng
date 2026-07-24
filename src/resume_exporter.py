"""
简历导出模块
负责把优化后的简历文本导出为 Word 文档
"""

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def extract_optimized_resume(analysis_text: str) -> str:
    """
    从 AI 分析结果中提取优化后的简历内容

    Args:
        analysis_text: AI 返回的完整分析文本

    Returns:
        优化后的简历文本
    """
    # 尝试匹配 ===OPTIMIZED_RESUME_START=== 标记
    match = re.search(
        r"===OPTIMIZED_RESUME_START===(.*?)===OPTIMIZED_RESUME_END===",
        analysis_text,
        re.DOTALL,
    )
    if match:
        return match.group(1).strip()

    # 兼容旧格式：匹配 "## 优化后的简历" 标题后的内容
    match = re.search(r"##\s*优化后的简历\s*\n(.*)", analysis_text, re.DOTALL)
    if match:
        return match.group(1).strip()

    # 如果没找到，返回空字符串
    return ""


def _set_cell_font(run, font_name="微软雅黑", font_size=Pt(10.5), bold=False, color=RGBColor(0, 0, 0)):
    """统一设置 run 的字体"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _parse_inline_formatting(paragraph, text: str):
    """
    解析行内 **加粗** 并写入段落
    """
    # 按 ** 分段
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        else:
            run.text = part
        _set_cell_font(run)


def create_resume_docx(resume_text: str) -> BytesIO:
    """
    把简历文本转换为 Word 文档

    Args:
        resume_text: 简历文本内容

    Returns:
        Word 文档的字节流
    """
    doc = Document()

    # 设置页面边距（窄边距，适合简历）
    sections = doc.sections[0]
    sections.top_margin = Cm(1.5)
    sections.bottom_margin = Cm(1.5)
    sections.left_margin = Cm(1.8)
    sections.right_margin = Cm(1.8)

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 按行处理，识别标题层级和列表
    lines = resume_text.split("\n")
    in_list = False

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line.strip():
            in_list = False
            continue

        stripped = line.strip()

        # 判断标题层级
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title_text = heading_match.group(2).strip()

            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            _parse_inline_formatting(p, title_text)
            # 设置标题格式
            for run in p.runs:
                run.bold = True
                if level == 1:
                    run.font.size = Pt(16)
                    p.space_before = Pt(14)
                    p.space_after = Pt(8)
                elif level == 2:
                    run.font.size = Pt(13)
                    p.space_before = Pt(12)
                    p.space_after = Pt(6)
                else:
                    run.font.size = Pt(11)
                    p.space_before = Pt(8)
                    p.space_after = Pt(4)
            in_list = False
            continue

        # 判断列表项
        list_match = re.match(r"^([\*\-\•])\s+(.+)$", stripped)
        if list_match:
            content = list_match.group(2).strip()
            p = doc.add_paragraph(style="List Bullet")
            _parse_inline_formatting(p, content)
            p.space_after = Pt(3)
            in_list = True
            continue

        # 普通段落
        p = doc.add_paragraph()
        _parse_inline_formatting(p, stripped)
        p.space_after = Pt(4)
        in_list = False

    # 保存到字节流
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_optimized_resume_to_docx(analysis_text: str) -> tuple[BytesIO, str]:
    """
    从分析结果中提取优化后的简历，并生成 Word 文档

    Args:
        analysis_text: AI 返回的完整分析文本

    Returns:
        (Word 文档字节流, 优化后的简历文本)
    """
    resume_text = extract_optimized_resume(analysis_text)
    if not resume_text:
        raise ValueError("未能从分析结果中提取到优化后的简历")

    docx_buffer = create_resume_docx(resume_text)
    return docx_buffer, resume_text
